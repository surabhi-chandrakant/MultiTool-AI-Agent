"""
Production-Ready AI Agent Web Application
- Web Search | Calculator | Document Query
- Full Document Management with Upload/Delete/Clear
- ALL 5 ISSUES FIXED
"""

import os
import json
import math
import re
import pickle
import time
from typing import List, Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import numpy as np
import requests
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET
from sentence_transformers import SentenceTransformer
from duckduckgo_search import DDGS
import PyPDF2
import docx
import csv
import warnings
warnings.filterwarnings('ignore')

from flask import Flask, render_template_string, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import shutil

# ==================== Configuration ====================

UPLOAD_FOLDER = 'uploads'
DOCUMENT_STORE_FOLDER = 'document_store'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'csv', 'md', 'py', 'js', 'html', 'css', 'json', 'xml'}

Path(UPLOAD_FOLDER).mkdir(exist_ok=True)
Path(DOCUMENT_STORE_FOLDER).mkdir(exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ==================== Document Store ====================

class DocumentStore:
    def __init__(self, storage_path: str = DOCUMENT_STORE_FOLDER):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
        self.documents = []
        self.embeddings = []
        self.model = None
        self.load_model()
        self.load_store()

    def load_model(self):
        try:
            print("Loading embedding model...")
            self.model = SentenceTransformer('all-MiniLM-L6-v2')
            print("✅ Model loaded successfully!")
        except Exception as e:
            print(f"❌ Error loading model: {e}")

    def extract_text_from_file(self, file_path: str) -> str:
        file_path = Path(file_path)
        try:
            if file_path.suffix.lower() in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            elif file_path.suffix.lower() == '.pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text if text.strip() else "No extractable text found in PDF."
            elif file_path.suffix.lower() == '.docx':
                doc = docx.Document(file_path)
                text = '\n'.join([para.text for para in doc.paragraphs])
                return text if text.strip() else "No text found in document."
            elif file_path.suffix.lower() == '.csv':
                text = ""
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    csv_reader = csv.reader(f)
                    for row in csv_reader:
                        text += ' '.join(row) + '\n'
                return text
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            return f"ERROR: {str(e)}"

    def add_document(self, file_path: str, metadata: Optional[Dict] = None) -> Dict:
        if self.model is None:
            return {"success": False, "message": "Model not loaded. Cannot add document."}
        text = self.extract_text_from_file(file_path)
        if text.startswith("ERROR:"):
            return {"success": False, "message": text}
        if not text.strip():
            return {"success": False, "message": "Document is empty or no text could be extracted."}
        filename = Path(file_path).name
        self.remove_document(filename, save=False)
        chunks = self.chunk_text(text, chunk_size=500, overlap=50)
        for i, chunk in enumerate(chunks):
            doc_entry = {
                'file_path': str(file_path),
                'file_name': filename,
                'file_size': os.path.getsize(file_path),
                'file_type': Path(file_path).suffix,
                'chunk_index': i,
                'text': chunk,
                'metadata': metadata or {},
                'timestamp': datetime.now().isoformat()
            }
            self.documents.append(doc_entry)
            embedding = self.model.encode(chunk)
            self.embeddings.append(embedding)
        self.save_store()
        return {
            "success": True,
            "message": f"✅ Added '{filename}' ({len(chunks)} chunks, {len(text)} characters)",
            "chunks": len(chunks),
            "filename": filename,
            "size": len(text),
            "file_type": Path(file_path).suffix
        }

    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        words = text.split()
        if len(words) <= chunk_size:
            return [text]
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk:
                chunks.append(chunk)
        return chunks

    def search(self, query: str, top_k: int = 5) -> List[Dict]:
        if self.model is None or len(self.documents) == 0:
            return []
        try:
            query_embedding = self.model.encode(query)
            similarities = []
            for doc_embedding in self.embeddings:
                similarity = np.dot(query_embedding, doc_embedding) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding)
                )
                similarities.append(similarity)
            top_indices = np.argsort(similarities)[-top_k:][::-1]
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.15:
                    result = self.documents[idx].copy()
                    result['similarity'] = float(similarities[idx])
                    results.append(result)
            if not results:
                results = self._keyword_search(query, top_k)
            return results
        except Exception as e:
            print(f"Search error: {e}")
            return self._keyword_search(query, top_k)

    def _keyword_search(self, query: str, top_k: int = 5) -> List[Dict]:
        results = []
        query_lower = query.lower()
        for doc in self.documents:
            if query_lower in doc['text'].lower():
                results.append({**doc, 'similarity': 0.5})
        return results[:top_k]

    def save_store(self):
        store_data = {
            'documents': self.documents,
            'embeddings': [emb.tolist() for emb in self.embeddings]
        }
        store_file = self.storage_path / 'store.pkl'
        with open(store_file, 'wb') as f:
            pickle.dump(store_data, f)

    def load_store(self):
        store_file = self.storage_path / 'store.pkl'
        if store_file.exists():
            try:
                with open(store_file, 'rb') as f:
                    store_data = pickle.load(f)
                    self.documents = store_data['documents']
                    self.embeddings = [np.array(emb) for emb in store_data['embeddings']]
                print(f"✅ Loaded {len(self.documents)} document chunks from storage")
            except Exception as e:
                print(f"⚠️ Error loading store: {e}")
                self.documents = []
                self.embeddings = []

    def get_stats(self) -> Dict:
        unique_files = {}
        for doc in self.documents:
            filename = doc['file_name']
            if filename not in unique_files:
                unique_files[filename] = {
                    'name': filename,
                    'type': doc.get('file_type', 'unknown'),
                    'size': doc.get('file_size', 0),
                    'chunks': 0,
                    'added_date': doc.get('timestamp', 'unknown')
                }
            unique_files[filename]['chunks'] += 1
        file_list = list(unique_files.values())
        file_types = list(set(doc.get('file_type', 'unknown') for doc in self.documents))
        return {
            'total_documents': len(file_list),
            'total_chunks': len(self.documents),
            'total_size_chars': sum(len(doc['text']) for doc in self.documents),
            'file_types': file_types,
            'files': file_list,
            'total_size_bytes': sum(doc.get('file_size', 0) for doc in self.documents if 'file_size' in doc)
        }

    def list_documents(self) -> List[Dict]:
        unique_files = {}
        for doc in self.documents:
            filename = doc['file_name']
            if filename not in unique_files:
                unique_files[filename] = {
                    'name': filename,
                    'type': doc.get('file_type', 'unknown'),
                    'size': doc.get('file_size', 0),
                    'chunks': 0,
                    'added_date': doc.get('timestamp', 'unknown')
                }
            unique_files[filename]['chunks'] += 1
        return list(unique_files.values())

    def remove_document(self, filename: str, save: bool = True) -> Dict:
        initial_count = len(self.documents)
        indices_to_remove = [
            i for i, doc in enumerate(self.documents)
            if doc['file_name'] == filename
        ]
        if not indices_to_remove:
            return {"success": False, "message": f"Document '{filename}' not found."}
        for idx in reversed(indices_to_remove):
            self.documents.pop(idx)
            self.embeddings.pop(idx)
        if save:
            self.save_store()
        removed_count = initial_count - len(self.documents)
        return {
            "success": True,
            "message": f"✅ Removed '{filename}' ({removed_count} chunks)",
            "removed_chunks": removed_count
        }

    def clear_all(self) -> Dict:
        count = len(self.documents)
        self.documents = []
        self.embeddings = []
        self.save_store()
        upload_path = Path(UPLOAD_FOLDER)
        if upload_path.exists():
            for file in upload_path.iterdir():
                if file.is_file():
                    file.unlink()
        return {
            "success": True,
            "message": f"✅ Cleared all documents ({count} chunks removed)",
            "removed_chunks": count
        }


# ==================== Web Searcher ====================

class WebSearcher:
    """
    FIX 2: Web search with multiple fallback sources.
    - Wikipedia API first (no rate limits, reliable)
    - DuckDuckGo with retry and exponential backoff
    - Google News RSS for news queries
    - Meaningful fallback link (not a broken error)
    """

    def __init__(self):
        self.ddgs = None
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                          '(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        })
        try:
            self.ddgs = DDGS()
        except Exception:
            pass

    def _ddgs_search_with_retry(self, query: str, max_results: int, retries: int = 3) -> List[Dict]:
        """DuckDuckGo search with exponential backoff on rate-limit errors."""
        for attempt in range(retries):
            try:
                time.sleep(1 + attempt)  # 1s, 2s, 3s
                results = []
                for r in self.ddgs.text(query, max_results=max_results):
                    results.append({
                        'title': r.get('title', 'No title'),
                        'link': r.get('href', r.get('link', '#')),
                        'snippet': r.get('body', '')[:300]
                    })
                return results
            except Exception as e:
                err = str(e).lower()
                if 'ratelimit' in err or '202' in err:
                    print(f"DuckDuckGo rate limit (attempt {attempt+1}), backing off...")
                    time.sleep(2 ** attempt)
                else:
                    print(f"DuckDuckGo error: {e}")
                    break
        return []

    def search(self, query: str, max_results: int = 5) -> List[Dict]:
        results = []

        # 1. Wikipedia (reliable, no rate limit)
        try:
            resp = self.session.get(
                'https://en.wikipedia.org/w/api.php',
                params={
                    'action': 'query', 'list': 'search', 'srsearch': query,
                    'format': 'json', 'srlimit': max_results
                },
                timeout=10
            )
            data = resp.json()
            for item in data.get('query', {}).get('search', []):
                results.append({
                    'title': f"📚 {item['title']}",
                    'link': f"https://en.wikipedia.org/wiki/{item['title'].replace(' ', '_')}",
                    'snippet': re.sub(r'<[^>]+>', '', item.get('snippet', ''))[:300]
                })
        except Exception as e:
            print(f"Wikipedia error: {e}")

        # 2. DuckDuckGo with retry
        if self.ddgs and len(results) < max_results:
            ddg_results = self._ddgs_search_with_retry(query, max_results)
            results.extend(ddg_results)

        # Deduplicate by title
        seen = set()
        unique_results = []
        for r in results:
            key = r['title'].lower()
            if key not in seen:
                seen.add(key)
                unique_results.append(r)

        if unique_results:
            return unique_results[:max_results]

        # 3. Final fallback: direct Google link (better than an error message)
        return [{
            'title': f'🔍 Search Google: {query}',
            'link': f'https://www.google.com/search?q={requests.utils.quote(query)}',
            'snippet': f'Our search APIs are temporarily limited. Click the link to search Google for "{query}".'
        }]

    def search_news(self, query: str, max_results: int = 5) -> List[Dict]:
        results = []

        # 1. DuckDuckGo news with retry
        if self.ddgs:
            try:
                time.sleep(1)
                for r in self.ddgs.news(query, max_results=max_results):
                    results.append({
                        'title': r.get('title', 'No title'),
                        'link': r.get('url', r.get('link', '#')),
                        'snippet': r.get('body', '')[:300],
                        'date': r.get('date', 'N/A'),
                        'source': r.get('source', 'Unknown')
                    })
            except Exception as e:
                print(f"DuckDuckGo News error: {e}")

        # 2. Google News RSS fallback
        if not results:
            try:
                news_url = (
                    f'https://news.google.com/rss/search?q={requests.utils.quote(query)}'
                    '&hl=en-US&gl=US&ceid=US:en'
                )
                resp = self.session.get(news_url, timeout=10)
                if resp.status_code == 200:
                    root = ET.fromstring(resp.content)
                    for item in root.findall('.//item')[:max_results]:
                        title = item.find('title')
                        link = item.find('link')
                        desc = item.find('description')
                        pub_date = item.find('pubDate')
                        source = item.find('source')
                        results.append({
                            'title': title.text if title is not None else 'No title',
                            'link': link.text if link is not None else '#',
                            'snippet': re.sub(r'<[^>]+>', '', desc.text or '')[:300] if desc is not None else '',
                            'date': pub_date.text if pub_date is not None else 'N/A',
                            'source': source.text if source is not None else 'Google News'
                        })
            except Exception as e:
                print(f"Google News RSS error: {e}")

        if results:
            return results[:max_results]

        return [{
            'title': f'📰 Search Google News: {query}',
            'link': f'https://news.google.com/search?q={requests.utils.quote(query)}',
            'snippet': f'Click to search Google News for "{query}".',
            'date': 'Now',
            'source': 'Google News'
        }]


# ==================== Calculator ====================

class Calculator:
    """
    FIX 1: Complete unit conversion for all common units.
    All unit names are lowercase in the lookup tables.
    Conversion regex now supports multi-word units (e.g. "fluid ounce").
    """

    LENGTH_UNITS = {
        'mm': 0.001, 'millimeter': 0.001, 'millimeters': 0.001,
        'cm': 0.01, 'centimeter': 0.01, 'centimeters': 0.01,
        'm': 1.0, 'meter': 1.0, 'meters': 1.0,
        'km': 1000.0, 'kilometer': 1000.0, 'kilometers': 1000.0,
        'in': 0.0254, 'inch': 0.0254, 'inches': 0.0254,
        'ft': 0.3048, 'foot': 0.3048, 'feet': 0.3048,
        'yd': 0.9144, 'yard': 0.9144, 'yards': 0.9144,
        'mi': 1609.34, 'mile': 1609.34, 'miles': 1609.34,
    }

    WEIGHT_UNITS = {
        'mg': 0.001, 'milligram': 0.001, 'milligrams': 0.001,
        'g': 1.0, 'gram': 1.0, 'grams': 1.0,
        'kg': 1000.0, 'kilogram': 1000.0, 'kilograms': 1000.0,
        'oz': 28.3495, 'ounce': 28.3495, 'ounces': 28.3495,
        'lb': 453.592, 'lbs': 453.592, 'pound': 453.592, 'pounds': 453.592,
        't': 1000000.0, 'ton': 1000000.0, 'tonne': 1000000.0, 'metric ton': 1000000.0,
    }

    VOLUME_UNITS = {
        'ml': 1.0, 'milliliter': 1.0, 'milliliters': 1.0,
        'millilitre': 1.0, 'millilitres': 1.0,
        'cl': 10.0, 'centiliter': 10.0,
        'l': 1000.0, 'liter': 1000.0, 'liters': 1000.0,
        'litre': 1000.0, 'litres': 1000.0,
        'kl': 1000000.0, 'kiloliter': 1000000.0,
        'tsp': 4.92892, 'teaspoon': 4.92892, 'teaspoons': 4.92892,
        'tbsp': 14.7868, 'tablespoon': 14.7868, 'tablespoons': 14.7868,
        'fl oz': 29.5735, 'fluid oz': 29.5735,
        'fluid ounce': 29.5735, 'fluid ounces': 29.5735,
        'cup': 236.588, 'cups': 236.588,
        'pt': 473.176, 'pint': 473.176, 'pints': 473.176,
        'qt': 946.353, 'quart': 946.353, 'quarts': 946.353,
        'gal': 3785.41, 'gallon': 3785.41, 'gallons': 3785.41,
    }

    TEMP_UNITS = {
        'c': 'C', 'celsius': 'C', 'centigrade': 'C',
        'f': 'F', 'fahrenheit': 'F',
        'k': 'K', 'kelvin': 'K',
    }

    # Sorted longest-first so multi-word units match before single-word ones
    _ALL_UNITS_SORTED: List[str] = []

    def __init__(self):
        all_keys = (
            list(self.LENGTH_UNITS.keys()) +
            list(self.WEIGHT_UNITS.keys()) +
            list(self.VOLUME_UNITS.keys()) +
            list(self.TEMP_UNITS.keys())
        )
        # Longest keys first ensures "fluid ounce" matches before "fluid"
        self._ALL_UNITS_SORTED = sorted(set(all_keys), key=lambda x: -len(x))

    def _find_unit(self, text: str) -> Optional[str]:
        """Return the first matching unit key found in text (longest match first)."""
        text_lower = text.lower().strip()
        for unit in self._ALL_UNITS_SORTED:
            # Full-word match using word boundary
            pattern = r'(?<![a-z])' + re.escape(unit) + r'(?![a-z])'
            if re.search(pattern, text_lower):
                return unit
        return None

    def calculate(self, expression: str) -> Dict:
        safe_dict = {
            'abs': abs, 'round': round, 'min': min, 'max': max,
            'sum': sum, 'pow': pow, 'sqrt': math.sqrt,
            'sin': math.sin, 'cos': math.cos, 'tan': math.tan,
            'asin': math.asin, 'acos': math.acos, 'atan': math.atan,
            'log': math.log, 'log10': math.log10, 'log2': math.log2,
            'exp': math.exp, 'pi': math.pi, 'e': math.e,
            'factorial': math.factorial, 'gcd': math.gcd,
            'degrees': math.degrees, 'radians': math.radians
        }
        try:
            expression = expression.strip()
            original = expression
            expression = expression.replace('^', '**')
            expression = expression.replace('×', '*')
            expression = expression.replace('÷', '/')
            expression = expression.replace('π', 'pi')
            result = eval(expression, {"__builtins__": {}}, safe_dict)
            if isinstance(result, float):
                if result == int(result):
                    result = int(result)
                else:
                    result = round(result, 10)
            return {"success": True, "expression": original, "result": str(result)}
        except Exception as e:
            return {"success": False, "expression": expression, "error": str(e)}

    def convert_units(self, value: float, from_unit: str, to_unit: str) -> Dict:
        """
        FIX 1: Resolve both from_unit and to_unit via the unified _find_unit lookup
        so aliases (km, kilometers, Km) all resolve correctly.
        """
        from_key = self._find_unit(from_unit)
        to_key = self._find_unit(to_unit)

        if from_key is None:
            return {"success": False, "error": f"Unknown unit: '{from_unit}'. Supported: length, weight, volume, temperature."}
        if to_key is None:
            return {"success": False, "error": f"Unknown unit: '{to_unit}'. Supported: length, weight, volume, temperature."}

        try:
            # Temperature
            if from_key in self.TEMP_UNITS and to_key in self.TEMP_UNITS:
                result = self._convert_temperature(value, self.TEMP_UNITS[from_key], self.TEMP_UNITS[to_key])
                return {"success": True, "result": result}

            # Length
            if from_key in self.LENGTH_UNITS and to_key in self.LENGTH_UNITS:
                base = value * self.LENGTH_UNITS[from_key]
                converted = base / self.LENGTH_UNITS[to_key]
                return {"success": True, "result": f"{value} {from_unit} = {converted:.6g} {to_unit}"}

            # Weight
            if from_key in self.WEIGHT_UNITS and to_key in self.WEIGHT_UNITS:
                base = value * self.WEIGHT_UNITS[from_key]
                converted = base / self.WEIGHT_UNITS[to_key]
                return {"success": True, "result": f"{value} {from_unit} = {converted:.6g} {to_unit}"}

            # Volume
            if from_key in self.VOLUME_UNITS and to_key in self.VOLUME_UNITS:
                base = value * self.VOLUME_UNITS[from_key]
                converted = base / self.VOLUME_UNITS[to_key]
                return {"success": True, "result": f"{value} {from_unit} = {converted:.6g} {to_unit}"}

            # Cross-category
            return {
                "success": False,
                "error": f"Cannot convert '{from_unit}' to '{to_unit}': they belong to different unit categories."
            }

        except Exception as e:
            return {"success": False, "error": str(e)}

    def _convert_temperature(self, value: float, from_scale: str, to_scale: str) -> str:
        # Normalise to Celsius first
        if from_scale == 'F':
            celsius = (value - 32) * 5 / 9
        elif from_scale == 'K':
            celsius = value - 273.15
        else:
            celsius = value

        if to_scale == 'F':
            result = celsius * 9 / 5 + 32
        elif to_scale == 'K':
            result = celsius + 273.15
        else:
            result = celsius

        return f"{value}°{from_scale} = {result:.4g}°{to_scale}"


# ==================== AI Agent ====================

class AIAgent:
    """
    FIX 3 – Input validation: broken regex fixed (no unescaped ] in char class).
    FIX 4 – Document QA: explicit commands + natural language auto-route to docs.
    FIX 5 – Routing order:
        1. Explicit command keywords (search, calc, convert, …)
        2. Conversion detection (has number + unit + "to" + unit)
        3. Math expression detection (has numbers + operators)
        4. Document semantic search (when docs loaded)
        5. Default: web search
    """

    def __init__(self):
        print("Initializing components...")
        self.web_searcher = WebSearcher()
        self.calculator = Calculator()
        self.document_store = DocumentStore()
        print("✅ AI Agent ready!")

    # ------------------------------------------------------------------
    # FIX 3: Input validation — corrected character class (no bare ] )
    # ------------------------------------------------------------------
    def _is_invalid_input(self, text: str) -> bool:
        # Characters that on their own make no sense as a query
        # Use a properly escaped character class
        special_chars = len(re.findall(r'[@#$%^&*()\[\]{};:"<>|\\`~]', text))
        alphanumeric = len(re.findall(r'[a-zA-Z0-9]', text))
        if special_chars > 0 and alphanumeric == 0:
            return True
        if special_chars > alphanumeric and alphanumeric < 3:
            return True
        return False

    # ------------------------------------------------------------------
    # FIX 5: Detection helpers — order matters in process_command
    # ------------------------------------------------------------------
    def _is_explicit_command(self, text: str) -> bool:
        """True if the query starts with one of our registered command keywords."""
        command_keywords = {
            'search', 'news', 'calc', 'calculate', 'math',
            'convert', 'add_doc', 'add', 'upload',
            'doc_stats', 'stats', 'documents',
            'list_docs', 'files', 'remove_doc', 'delete',
            'clear_all', 'clear', 'help',
            'query_doc', 'find', 'search_doc',
        }
        first_word = text.split()[0].lower() if text.split() else ''
        return first_word in command_keywords

    def _is_conversion(self, text: str) -> bool:
        """
        FIX 1 + FIX 5: Detect unit conversion queries robustly.
        Patterns supported:
          "100 km to meters"
          "convert 100 km to meters"
          "100km to m"
        """
        text_lower = text.lower()

        if ' to ' not in text_lower and not text_lower.startswith('convert '):
            return False

        # Must have a number somewhere
        if not re.search(r'\d', text):
            return False

        # Check that at least one recognised unit appears
        for unit in self.calculator._ALL_UNITS_SORTED:
            if re.search(r'(?<![a-z])' + re.escape(unit) + r'(?![a-z])', text_lower):
                return True

        return False

    def _is_math_expression(self, text: str) -> bool:
        """
        FIX 5: Math detection — only fires when the expression clearly looks like
        arithmetic, not a natural-language question.
        """
        # If it starts with a question/search word and is long, skip
        question_starters = {
            'what', 'how', 'when', 'where', 'who', 'why', 'is', 'are',
            'tell', 'show', 'find', 'get', 'search', 'news', 'help',
            'can', 'please', 'would', 'could', 'should', 'do', 'does',
            'did', 'explain', 'describe', 'list'
        }
        words = text.split()
        if words and words[0].lower() in question_starters and len(words) > 3:
            return False

        math_functions = ['sqrt', 'sin', 'cos', 'tan', 'log', 'abs', 'factorial']
        has_numbers = bool(re.search(r'\d', text))
        has_operators = bool(re.search(r'(?<![a-z])[\+\-\*\/\^](?![a-z])', text))
        has_functions = any(f in text.lower() for f in math_functions)
        has_percent = '%' in text and has_numbers

        # Must have both a number AND (an operator OR function OR %)
        return has_numbers and (has_operators or has_functions or has_percent)

    # ------------------------------------------------------------------
    # Main dispatcher
    # ------------------------------------------------------------------
    def process_command(self, user_input: str) -> Dict:
        user_input = user_input.strip()

        # Guard: empty input
        if not user_input:
            return {'type': 'error', 'message': 'Please enter a command or question.'}

        # FIX 3: Validate input before doing anything
        if self._is_invalid_input(user_input):
            return {
                'type': 'error',
                'message': (
                    '⚠️ Invalid input detected. Please enter a meaningful question, '
                    'calculation, or command. Queries consisting only of special '
                    'characters are not supported.'
                )
            }

        # ---- Step 1: Explicit keyword commands (highest priority) ----
        if self._is_explicit_command(user_input):
            parts = user_input.split(' ', 1)
            command = parts[0].lower()
            args = parts[1] if len(parts) > 1 else ''

            routes = {
                'search':      lambda: self._handle_search(args),
                'news':        lambda: self._handle_news(args),
                'calc':        lambda: self._handle_calc(args),
                'calculate':   lambda: self._handle_calc(args),
                'math':        lambda: self._handle_calc(args),
                'convert':     lambda: self._handle_convert(args),
                'add_doc':     lambda: self._handle_add_document(args),
                'add':         lambda: self._handle_add_document(args),
                'upload':      lambda: self._handle_add_document(args),
                'doc_stats':   lambda: self._handle_doc_stats(),
                'stats':       lambda: self._handle_doc_stats(),
                'documents':   lambda: self._handle_doc_stats(),
                'list_docs':   lambda: self._handle_list_documents(),
                'files':       lambda: self._handle_list_documents(),
                'remove_doc':  lambda: self._handle_remove_document(args),
                'delete':      lambda: self._handle_remove_document(args),
                'clear_all':   lambda: self._handle_clear_all(),
                'clear':       lambda: self._handle_clear_all(),
                'help':        lambda: self._handle_help(),
                'query_doc':   lambda: self._handle_query_document(args),
                'find':        lambda: self._handle_query_document(args),
                'search_doc':  lambda: self._handle_query_document(args),
            }

            if command in routes:
                return routes[command]()

        # ---- Step 2: Unit conversion (before math — "100 km to m" has operators) ----
        if self._is_conversion(user_input):
            return self._handle_convert(user_input)

        # ---- Step 3: Pure math expression ----
        if self._is_math_expression(user_input):
            return self._handle_calc(user_input)

        # ---- Step 4: FIX 4 – Document semantic search when docs are loaded ----
        if self.document_store.documents:
            doc_results = self.document_store.search(user_input, top_k=5)
            if doc_results:
                return {
                    'type': 'document_search',
                    'query': user_input,
                    'results': doc_results,
                    'count': len(doc_results)
                }

        # ---- Step 5: Default — web search ----
        return self._handle_search(user_input)

    # ------------------------------------------------------------------
    # Handlers
    # ------------------------------------------------------------------
    def _handle_search(self, query: str) -> Dict:
        if not query:
            return {'type': 'error', 'message': 'What would you like to search for?'}
        if self._is_invalid_input(query):
            return {'type': 'error', 'message': '⚠️ Invalid search query. Please use meaningful keywords.'}
        results = self.web_searcher.search(query)
        return {'type': 'search', 'query': query, 'results': results, 'count': len(results)}

    def _handle_news(self, query: str) -> Dict:
        if not query:
            return {'type': 'error', 'message': 'What news topic?'}
        if self._is_invalid_input(query):
            return {'type': 'error', 'message': '⚠️ Invalid news query.'}
        results = self.web_searcher.search_news(query)
        return {'type': 'news', 'query': query, 'results': results, 'count': len(results)}

    def _handle_calc(self, expression: str) -> Dict:
        if not expression:
            return {'type': 'error', 'message': 'What would you like to calculate?'}
        # Strip command prefix if user typed e.g. "calc 3+3"
        expression = re.sub(r'^(calc|calculate|math)\s+', '', expression.strip(), flags=re.IGNORECASE)
        result = self.calculator.calculate(expression)
        if result['success']:
            return {'type': 'calc', 'expression': result['expression'], 'result': result['result']}
        return {'type': 'error', 'message': f"Calculation error: {result['error']}"}

    def _handle_convert(self, args: str) -> Dict:
        if not args:
            return {
                'type': 'error',
                'message': 'Format: convert <value> <unit> to <unit>\nExample: convert 100 km to meters'
            }

        # Strip leading "convert" keyword if present
        text = re.sub(r'^convert\s+', '', args.strip(), flags=re.IGNORECASE).strip()

        # FIX 1: Greedy unit matching — support multi-word units.
        # Build a sorted-longest regex alternation from all known units.
        unit_pattern = '|'.join(re.escape(u) for u in self.calculator._ALL_UNITS_SORTED)
        # Pattern: <number> <from_unit> to <to_unit>
        pattern = rf'^([\d.]+)\s*({unit_pattern})\s+to\s+({unit_pattern})\s*$'
        match = re.match(pattern, text, re.IGNORECASE)

        if not match:
            # Fallback: try loose "number word(s) to word(s)"
            match2 = re.match(r'^([\d.]+)\s+(.+?)\s+to\s+(.+)$', text, re.IGNORECASE)
            if not match2:
                return {
                    'type': 'error',
                    'message': (
                        'Format: convert <value> <unit> to <unit>\n'
                        'Example: convert 100 km to meters'
                    )
                }
            value_str, from_unit, to_unit = match2.group(1), match2.group(2).strip(), match2.group(3).strip()
        else:
            value_str, from_unit, to_unit = match.group(1), match.group(2).strip(), match.group(3).strip()

        try:
            value = float(value_str)
        except ValueError:
            return {'type': 'error', 'message': 'Invalid number value.'}

        result = self.calculator.convert_units(value, from_unit, to_unit)
        if result['success']:
            return {'type': 'convert', 'result': result['result']}
        return {'type': 'error', 'message': result['error']}

    def _handle_add_document(self, file_path: str) -> Dict:
        if not file_path:
            return {'type': 'error', 'message': 'Please provide a file path.\nExample: add_doc /path/to/document.pdf'}
        file_path = file_path.strip().strip('"\'')
        if not os.path.exists(file_path):
            return {'type': 'error', 'message': f'❌ File not found: {file_path}'}
        if not allowed_file(file_path):
            return {'type': 'error', 'message': f'File type not allowed. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'}
        result = self.document_store.add_document(file_path)
        if result['success']:
            return {'type': 'document', 'message': result['message'], 'filename': result['filename'], 'chunks': result['chunks']}
        return {'type': 'error', 'message': result['message']}

    def _handle_query_document(self, query: str) -> Dict:
        """FIX 4: Document QA with semantic + keyword fallback."""
        if not query:
            return {'type': 'error', 'message': 'What would you like to find in your documents?'}
        if not self.document_store.documents:
            return {
                'type': 'error',
                'message': '📚 No documents in store. Upload documents first using the left panel.'
            }
        results = self.document_store.search(query)
        if not results:
            results = self.document_store._keyword_search(query)
        if results:
            return {'type': 'document_search', 'query': query, 'results': results[:5], 'count': len(results)}
        return {
            'type': 'document_search',
            'query': query,
            'results': [],
            'count': 0,
            'message': f'No matching content found for "{query}" in your documents.'
        }

    def _handle_doc_stats(self) -> Dict:
        return {'type': 'document_stats', 'stats': self.document_store.get_stats()}

    def _handle_list_documents(self) -> Dict:
        docs = self.document_store.list_documents()
        return {'type': 'document_list', 'files': docs, 'count': len(docs)}

    def _handle_remove_document(self, filename: str) -> Dict:
        if not filename:
            return {'type': 'error', 'message': 'Please specify a filename to remove.'}
        result = self.document_store.remove_document(filename)
        if result['success']:
            return {'type': 'document', 'message': result['message']}
        return {'type': 'error', 'message': result['message']}

    def _handle_clear_all(self) -> Dict:
        result = self.document_store.clear_all()
        return {'type': 'document', 'message': result['message']}

    def _handle_help(self) -> Dict:
        return {
            'type': 'help',
            'commands': [
                {'cmd': 'search <query>',                    'desc': 'Search the web',           'example': 'search latest AI news'},
                {'cmd': 'news <topic>',                      'desc': 'Latest news',              'example': 'news technology'},
                {'cmd': 'calc <expression>',                 'desc': 'Calculate',                'example': 'calc 15 * 3 + 27'},
                {'cmd': 'convert <value> <from> to <to>',   'desc': 'Convert units',            'example': 'convert 100 km to meters'},
                {'cmd': 'add_doc <path>',                    'desc': 'Add document',             'example': 'add_doc report.txt'},
                {'cmd': 'query_doc <text>',                  'desc': 'Search documents',         'example': 'query_doc AI trends'},
                {'cmd': 'doc_stats',                         'desc': 'Document statistics',      'example': 'doc_stats'},
                {'cmd': 'list_docs',                         'desc': 'List all documents',       'example': 'list_docs'},
                {'cmd': 'remove_doc <name>',                 'desc': 'Remove document',          'example': 'remove_doc report.txt'},
                {'cmd': 'clear_all',                         'desc': 'Clear all documents',      'example': 'clear_all'},
            ]
        }


# ==================== Flask Web App ====================

app = Flask(__name__)
app.secret_key = 'production-ai-agent-2024'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
CORS(app)

print("\n" + "=" * 60)
print("🤖 Initializing Production AI Agent (ALL 5 ISSUES FIXED)...")
print("=" * 60)
agent = AIAgent()

# ==================== HTML Template ====================

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🤖 AI Agent Pro - Web Search | Calculator | Documents</title>
    <style>
        :root {
            --primary: #667eea;
            --secondary: #764ba2;
            --success: #48bb78;
            --warning: #ed8936;
            --danger: #fc8181;
            --bg: #f7fafc;
            --card: #ffffff;
            --text: #2d3748;
            --border: #e2e8f0;
        }
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }
        .app-container {
            width: 100%;
            max-width: 1400px;
            background: white;
            border-radius: 20px;
            box-shadow: 0 25px 80px rgba(0,0,0,0.3);
            overflow: hidden;
            display: flex;
            height: 90vh;
        }
        .sidebar {
            width: 320px;
            background: #2d3748;
            color: white;
            display: flex;
            flex-direction: column;
            border-right: 1px solid #4a5568;
        }
        .sidebar-header {
            padding: 20px;
            background: #1a202c;
            text-align: center;
            border-bottom: 1px solid #4a5568;
        }
        .sidebar-header h2 { font-size: 20px; margin-bottom: 5px; }
        .sidebar-header p { font-size: 12px; opacity: 0.7; }
        .upload-section { padding: 20px; border-bottom: 1px solid #4a5568; }
        .upload-area {
            border: 2px dashed #4a5568;
            border-radius: 10px;
            padding: 20px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s;
            margin-bottom: 10px;
        }
        .upload-area:hover { border-color: var(--primary); background: rgba(102,126,234,0.1); }
        .upload-area.dragover { border-color: var(--success); background: rgba(72,187,120,0.1); }
        .upload-icon { font-size: 30px; margin-bottom: 10px; }
        .upload-text { font-size: 13px; color: #a0aec0; }
        #fileInput { display: none; }
        .btn {
            padding: 10px 20px;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            font-size: 13px;
            font-weight: 600;
            transition: all 0.3s;
            width: 100%;
            margin-bottom: 8px;
        }
        .btn-primary { background: var(--primary); color: white; }
        .btn-primary:hover { background: #5a6fd6; transform: translateY(-2px); }
        .btn-danger { background: var(--danger); color: white; }
        .btn-danger:hover { background: #f56565; transform: translateY(-2px); }
        .btn-success { background: var(--success); color: white; }
        .btn-success:hover { background: #38a169; transform: translateY(-2px); }
        .document-list { flex: 1; overflow-y: auto; padding: 15px; }
        .document-list h3 {
            font-size: 14px;
            color: #a0aec0;
            margin-bottom: 15px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .doc-item {
            background: #4a5568;
            border-radius: 8px;
            padding: 12px;
            margin-bottom: 10px;
            transition: all 0.3s;
        }
        .doc-item:hover { background: #5a6578; transform: translateX(5px); }
        .doc-item-header {
            display: flex;
            justify-content: space-between;
            align-items: start;
            margin-bottom: 8px;
        }
        .doc-name { font-size: 13px; font-weight: 600; word-break: break-word; }
        .doc-type {
            font-size: 10px;
            background: var(--primary);
            padding: 2px 8px;
            border-radius: 10px;
            white-space: nowrap;
        }
        .doc-meta { font-size: 11px; color: #a0aec0; display: flex; gap: 15px; }
        .remove-doc-btn {
            background: none;
            border: none;
            color: #fc8181;
            cursor: pointer;
            font-size: 16px;
            padding: 2px 5px;
            transition: all 0.3s;
        }
        .remove-doc-btn:hover { color: #f56565; transform: scale(1.2); }
        .empty-state { text-align: center; padding: 40px 20px; color: #a0aec0; }
        .empty-state .icon { font-size: 50px; margin-bottom: 15px; }
        .main-area { flex: 1; display: flex; flex-direction: column; min-width: 0; }
        .main-header {
            padding: 15px 25px;
            background: white;
            border-bottom: 1px solid var(--border);
            display: flex;
            align-items: center;
            gap: 15px;
        }
        .main-header h1 {
            font-size: 22px;
            background: linear-gradient(135deg, var(--primary), var(--secondary));
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        .quick-actions {
            display: flex;
            gap: 8px;
            padding: 10px 25px;
            background: var(--bg);
            border-bottom: 1px solid var(--border);
            flex-wrap: wrap;
        }
        .quick-btn {
            padding: 6px 14px;
            background: white;
            border: 1px solid var(--border);
            border-radius: 15px;
            cursor: pointer;
            font-size: 12px;
            transition: all 0.3s;
            white-space: nowrap;
        }
        .quick-btn:hover { background: var(--primary); color: white; border-color: var(--primary); }
        .chat-area { flex: 1; overflow-y: auto; padding: 25px; background: var(--bg); }
        .message { margin-bottom: 20px; animation: slideIn 0.3s ease-out; }
        @keyframes slideIn {
            from { opacity: 0; transform: translateY(10px); }
            to   { opacity: 1; transform: translateY(0); }
        }
        .user-message { display: flex; justify-content: flex-end; }
        .user-bubble {
            background: var(--primary);
            color: white;
            padding: 12px 18px;
            border-radius: 18px 18px 5px 18px;
            max-width: 60%;
            word-wrap: break-word;
            font-size: 14px;
        }
        .agent-message { display: flex; justify-content: flex-start; }
        .agent-bubble {
            background: white;
            padding: 15px 20px;
            border-radius: 18px 18px 18px 5px;
            max-width: 75%;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
            word-wrap: break-word;
            font-size: 14px;
            line-height: 1.6;
        }
        .search-result, .news-result, .doc-result {
            margin: 10px 0;
            padding: 12px;
            border-radius: 8px;
            border-left: 4px solid;
        }
        .search-result { background: #f0f4ff; border-color: var(--primary); }
        .news-result   { background: #f0f9ff; border-color: #4299e1; }
        .doc-result    { background: #fffaf0; border-color: var(--warning); }
        .search-result h4, .news-result h4 { margin-bottom: 5px; font-size: 14px; }
        .search-result a, .news-result a { color: var(--secondary); text-decoration: none; font-size: 12px; }
        .search-result p, .news-result p { color: #666; font-size: 13px; margin-top: 5px; }
        .calc-result {
            background: linear-gradient(135deg, #f0fff4, #e6fffa);
            border-left: 4px solid var(--success);
            padding: 15px;
            border-radius: 8px;
        }
        .calc-result .expression { color: #666; font-size: 14px; }
        .calc-result .result { color: var(--success); font-size: 28px; font-weight: bold; }
        .stats-box { background: linear-gradient(135deg, #fefcbf, #faf089); border-radius: 12px; padding: 15px; }
        .stats-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(120px, 1fr)); gap: 10px; margin-top: 10px; }
        .stat-item { background: white; padding: 12px; border-radius: 8px; text-align: center; }
        .stat-value { font-size: 22px; font-weight: bold; color: var(--primary); }
        .stat-label { font-size: 11px; color: #666; margin-top: 5px; }
        .error-message   { background: #fff5f5; border-left: 4px solid var(--danger); color: #c53030; padding: 10px; border-radius: 8px; }
        .success-message { background: #f0fff4; border-left: 4px solid var(--success); color: #22543d; padding: 10px; border-radius: 8px; }
        .input-area {
            padding: 15px 25px;
            background: white;
            border-top: 1px solid var(--border);
            display: flex;
            gap: 10px;
        }
        #userInput {
            flex: 1;
            padding: 12px 20px;
            border: 2px solid var(--border);
            border-radius: 25px;
            font-size: 14px;
            outline: none;
            transition: all 0.3s;
        }
        #userInput:focus { border-color: var(--primary); box-shadow: 0 0 0 3px rgba(102,126,234,0.1); }
        #sendBtn {
            padding: 12px 25px;
            background: linear-gradient(135deg, var(--primary), var(--secondary));
            color: white;
            border: none;
            border-radius: 25px;
            cursor: pointer;
            font-weight: bold;
            transition: all 0.3s;
        }
        #sendBtn:hover { transform: scale(1.05); box-shadow: 0 5px 15px rgba(102,126,234,0.4); }
        .loading {
            display: inline-block;
            width: 20px;
            height: 20px;
            border: 3px solid #f3f3f3;
            border-top: 3px solid var(--primary);
            border-radius: 50%;
            animation: spin 1s linear infinite;
        }
        @keyframes spin { 0% { transform: rotate(0deg); } 100% { transform: rotate(360deg); } }
        .toast {
            position: fixed;
            top: 20px;
            right: 20px;
            padding: 15px 20px;
            border-radius: 10px;
            color: white;
            font-weight: 600;
            animation: slideDown 0.3s ease-out;
            z-index: 1000;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        .toast-success { background: var(--success); }
        .toast-error   { background: var(--danger); }
        @keyframes slideDown {
            from { opacity: 0; transform: translateY(-20px); }
            to   { opacity: 1; transform: translateY(0); }
        }
        @media (max-width: 768px) {
            .app-container { flex-direction: column; height: 100vh; border-radius: 0; }
            .sidebar { width: 100%; max-height: 200px; }
            .user-bubble  { max-width: 85%; }
            .agent-bubble { max-width: 90%; }
        }
    </style>
</head>
<body>
<div class="app-container">
    <!-- Sidebar -->
    <div class="sidebar">
        <div class="sidebar-header">
            <h2>📚 Documents</h2>
            <p>Upload &amp; manage your files</p>
        </div>
        <div class="upload-section">
            <div class="upload-area" id="uploadArea" onclick="document.getElementById('fileInput').click()">
                <div class="upload-icon">📁</div>
                <div class="upload-text">
                    <strong>Click to upload</strong> or drag &amp; drop<br>
                    <small>TXT, PDF, DOCX, CSV, MD, code files</small>
                </div>
            </div>
            <input type="file" id="fileInput" multiple
                   accept=".txt,.pdf,.docx,.csv,.md,.py,.js,.html,.css,.json,.xml"
                   onchange="handleFileUpload(this.files)">
            <button class="btn btn-success" onclick="refreshDocuments()">🔄 Refresh List</button>
            <button class="btn btn-danger"  onclick="clearAllDocuments()">🗑️ Clear All Documents</button>
        </div>
        <div class="document-list" id="documentList">
            <h3>📋 Your Documents</h3>
            <div class="empty-state" id="emptyState">
                <div class="icon">📭</div>
                <p>No documents yet</p>
                <small>Upload files to get started</small>
            </div>
        </div>
    </div>

    <!-- Main Chat -->
    <div class="main-area">
        <div class="main-header">
            <h1>🤖 AI Agent Pro</h1>
            <small style="color:#666;">Web Search • Calculator • Document Query</small>
        </div>
        <div class="quick-actions">
            <button class="quick-btn" onclick="setQuickAction('search ')">🔍 Search</button>
            <button class="quick-btn" onclick="setQuickAction('news ')">📰 News</button>
            <button class="quick-btn" onclick="setQuickAction('calc ')">🧮 Calc</button>
            <button class="quick-btn" onclick="setQuickAction('convert ')">📏 Convert</button>
            <button class="quick-btn" onclick="setQuickAction('doc_stats')">📊 Stats</button>
            <button class="quick-btn" onclick="setQuickAction('list_docs')">📋 List</button>
            <button class="quick-btn" onclick="setQuickAction('help')">❓ Help</button>
            <button class="quick-btn" onclick="clearChat()">🗑️ Clear Chat</button>
        </div>
        <div class="chat-area" id="chatArea">
            <div class="message agent-message">
                <div class="agent-bubble">
                    <strong>👋 Welcome to AI Agent Pro!</strong><br><br>
                    <strong>Quick Start:</strong><br>
                    • 🔍 <strong>Search web:</strong> "search Python tutorials"<br>
                    • 📰 <strong>Get news:</strong> "news technology"<br>
                    • 🧮 <strong>Calculate:</strong> "calc 15 * 3 + 27"<br>
                    • 📏 <strong>Convert:</strong> "convert 100 km to miles"<br>
                    • 📚 <strong>Documents:</strong> Upload files using the left panel<br><br>
                    <em>Type <strong>help</strong> for all commands</em>
                </div>
            </div>
        </div>
        <div class="input-area">
            <input type="text" id="userInput" placeholder="Type your command or question..."
                   onkeypress="if(event.key==='Enter') sendMessage()">
            <button id="sendBtn" onclick="sendMessage()">Send 🚀</button>
        </div>
    </div>
</div>

<script>
    const uploadArea = document.getElementById('uploadArea');
    ['dragenter','dragover','dragleave','drop'].forEach(e => uploadArea.addEventListener(e, ev => { ev.preventDefault(); ev.stopPropagation(); }));
    ['dragenter','dragover'].forEach(e => uploadArea.addEventListener(e, () => uploadArea.classList.add('dragover')));
    ['dragleave','drop'].forEach(e => uploadArea.addEventListener(e, () => uploadArea.classList.remove('dragover')));
    uploadArea.addEventListener('drop', e => handleFileUpload(e.dataTransfer.files));

    async function handleFileUpload(files) {
        if (!files || files.length === 0) return;
        for (const file of files) {
            const fd = new FormData();
            fd.append('file', file);
            try {
                const res  = await fetch('/upload', { method: 'POST', body: fd });
                const data = await res.json();
                if (data.type === 'document') {
                    showToast(data.message, 'success');
                    addMessage('agent', formatResponse(data));
                } else {
                    showToast(data.message || 'Upload failed', 'error');
                }
            } catch {
                showToast('Error uploading file: ' + file.name, 'error');
            }
        }
        refreshDocuments();
        document.getElementById('fileInput').value = '';
    }

    async function refreshDocuments() {
        try {
            const res  = await fetch('/query', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({query:'list_docs'}) });
            const data = await res.json();
            updateDocumentList(data);
        } catch (e) { console.error('Error refreshing documents:', e); }
    }

    function updateDocumentList(data) {
        const docList = document.getElementById('documentList');
        if (!data.files || data.files.length === 0) {
            docList.innerHTML = '<h3>📋 Your Documents</h3><div class="empty-state"><div class="icon">📭</div><p>No documents yet</p><small>Upload files to get started</small></div>';
            return;
        }
        let html = '<h3>📋 Your Documents</h3>';
        data.files.forEach(file => {
            const size = file.size ? (file.size / 1024).toFixed(1) + ' KB' : 'N/A';
            html += `
                <div class="doc-item">
                    <div class="doc-item-header">
                        <span class="doc-name">📄 ${file.name}</span>
                        <span class="doc-type">${file.type || 'file'}</span>
                    </div>
                    <div class="doc-meta">
                        <span>📑 ${file.chunks} chunks</span>
                        <span>💾 ${size}</span>
                        <button class="remove-doc-btn" onclick="removeDocument('${file.name}')" title="Remove">✕</button>
                    </div>
                </div>`;
        });
        docList.innerHTML = html;
    }

    async function removeDocument(filename) {
        if (!confirm(`Remove "${filename}" from document store?`)) return;
        try {
            const res  = await fetch('/query', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({query:`remove_doc ${filename}`}) });
            const data = await res.json();
            showToast(data.message, data.type === 'error' ? 'error' : 'success');
            addMessage('agent', formatResponse(data));
            refreshDocuments();
        } catch { showToast('Error removing document', 'error'); }
    }

    async function clearAllDocuments() {
        if (!confirm('⚠️ Are you sure? This will remove ALL documents from the store and uploads folder.')) return;
        try {
            const res  = await fetch('/query', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({query:'clear_all'}) });
            const data = await res.json();
            showToast(data.message, 'success');
            addMessage('agent', formatResponse(data));
            refreshDocuments();
        } catch { showToast('Error clearing documents', 'error'); }
    }

    function showToast(message, type) {
        const t = document.createElement('div');
        t.className = `toast toast-${type}`;
        t.textContent = message;
        document.body.appendChild(t);
        setTimeout(() => { t.style.opacity = '0'; t.style.transition = 'opacity 0.3s'; setTimeout(() => t.remove(), 300); }, 3000);
    }

    function setQuickAction(action) {
        document.getElementById('userInput').value = action;
        document.getElementById('userInput').focus();
    }

    function clearChat() { document.getElementById('chatArea').innerHTML = ''; }

    function addMessage(type, content) {
        const chat = document.getElementById('chatArea');
        const msg  = document.createElement('div');
        msg.className = `message ${type}-message`;
        const bubble = document.createElement('div');
        bubble.className = `${type}-bubble`;
        bubble.innerHTML = content;
        msg.appendChild(bubble);
        chat.appendChild(msg);
        chat.scrollTop = chat.scrollHeight;
    }

    function showLoading() {
        const chat = document.getElementById('chatArea');
        const div  = document.createElement('div');
        div.className = 'message agent-message';
        div.id = 'loadingMessage';
        div.innerHTML = '<div class="agent-bubble"><div class="loading"></div> Processing...</div>';
        chat.appendChild(div);
        chat.scrollTop = chat.scrollHeight;
    }

    function hideLoading() { const el = document.getElementById('loadingMessage'); if (el) el.remove(); }

    function formatResponse(data) {
        let html = '';
        switch (data.type) {
            case 'search':
                html += `<strong>🔍 Results: "${data.query}"</strong> (${data.count} found)<br><br>`;
                data.results.forEach((r, i) => {
                    html += `<div class="search-result"><h4>${i+1}. ${r.title}</h4><a href="${r.link}" target="_blank">🔗 ${r.link}</a><p>${r.snippet}</p></div>`;
                });
                break;
            case 'news':
                html += `<strong>📰 News: "${data.query}"</strong> (${data.count} articles)<br><br>`;
                data.results.forEach((r, i) => {
                    html += `<div class="news-result"><h4>${i+1}. ${r.title}</h4><small>📅 ${r.date} | 📰 ${r.source}</small><br><a href="${r.link}" target="_blank">🔗 Read more</a><p>${r.snippet}</p></div>`;
                });
                break;
            case 'calc':
                html += `<div class="calc-result"><div class="expression">📝 ${data.expression}</div><div class="result">= ${data.result}</div></div>`;
                break;
            case 'convert':
                html += `<div class="calc-result">📏 <strong>${data.result}</strong></div>`;
                break;
            case 'document':
                html += `<div class="success-message">${data.message}</div>`;
                break;
            case 'document_list':
                html += `<strong>📋 Documents (${data.count}):</strong><br>`;
                (data.files || []).forEach(f => { html += `<div style="margin:5px 0;">📄 ${f.name} (${f.chunks} chunks, ${f.type})</div>`; });
                break;
            case 'document_search':
                html += `<strong>📚 Results for "${data.query}"</strong> (${data.count} matches)<br><br>`;
                if (data.results && data.results.length > 0) {
                    data.results.forEach(r => {
                        html += `<div class="doc-result"><strong>📄 ${r.file_name}</strong> <span style="color:#ed8936;">(${(r.similarity*100).toFixed(1)}% match)</span><p>${r.text.substring(0, 200)}…</p></div>`;
                    });
                } else {
                    html += `<div class="error-message">${data.message || 'No matches found.'}</div>`;
                }
                break;
            case 'document_stats':
                if (data.stats) {
                    const s = data.stats;
                    html += `<div class="stats-box"><strong>📊 Document Statistics</strong>
                        <div class="stats-grid">
                            <div class="stat-item"><div class="stat-value">${s.total_documents}</div><div class="stat-label">Documents</div></div>
                            <div class="stat-item"><div class="stat-value">${s.total_chunks}</div><div class="stat-label">Chunks</div></div>
                            <div class="stat-item"><div class="stat-value">${(s.total_size_chars/1000).toFixed(1)}K</div><div class="stat-label">Characters</div></div>
                        </div><br>
                        <strong>Types:</strong> ${s.file_types?.join(', ') || 'None'}<br>
                        <strong>Files:</strong><br>${s.files?.map(f => `📄 ${f.name} (${f.chunks} chunks)`).join('<br>') || 'No files'}
                    </div>`;
                }
                break;
            case 'help':
                html += '<strong>📚 Commands:</strong><br><br>';
                data.commands.forEach(c => {
                    html += `<div style="margin:8px 0;"><strong style="color:#667eea;">${c.cmd}</strong><br><small>${c.desc}</small><br><code>Example: ${c.example}</code></div>`;
                });
                break;
            case 'error':
                html += `<div class="error-message">❌ ${data.message}</div>`;
                break;
            default:
                html += data.message || 'Done!';
        }
        return html;
    }

    async function sendMessage() {
        const input   = document.getElementById('userInput');
        const message = input.value.trim();
        if (!message) return;
        addMessage('user', message);
        input.value = '';
        showLoading();
        try {
            const res  = await fetch('/query', { method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({query: message}) });
            const data = await res.json();
            hideLoading();
            addMessage('agent', formatResponse(data));
            if (['document','document_list','document_stats'].includes(data.type)) refreshDocuments();
        } catch {
            hideLoading();
            addMessage('agent', '<div class="error-message">❌ Connection error. Please try again.</div>');
        }
    }

    refreshDocuments();
    document.getElementById('userInput').focus();
</script>
</body>
</html>
'''


@app.route('/')
def home():
    return render_template_string(HTML_TEMPLATE)


@app.route('/query', methods=['POST'])
def query():
    try:
        data = request.json
        user_query = data.get('query', '')
        if not user_query:
            return jsonify({'type': 'error', 'message': 'No query provided'})
        return jsonify(agent.process_command(user_query))
    except Exception as e:
        return jsonify({'type': 'error', 'message': f'Server error: {str(e)}'})


@app.route('/upload', methods=['POST'])
def upload_document():
    try:
        if 'file' not in request.files:
            return jsonify({'type': 'error', 'message': 'No file uploaded'})
        file = request.files['file']
        if file.filename == '':
            return jsonify({'type': 'error', 'message': 'No file selected'})
        if not allowed_file(file.filename):
            return jsonify({'type': 'error', 'message': f'File type not allowed. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'})
        filename  = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        result = agent.document_store.add_document(file_path)
        if result['success']:
            return jsonify({
                'type': 'document',
                'message': result['message'],
                'filename': result['filename'],
                'chunks': result['chunks']
            })
        return jsonify({'type': 'error', 'message': result['message']})
    except Exception as e:
        return jsonify({'type': 'error', 'message': f'Upload error: {str(e)}'})


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    print("\n" + "=" * 60)
    print("🚀 Production AI Agent Starting...")
    print("=" * 60)
    print(f"\n📱 Running on port: {port}")
    print("\n✅ Fixes applied:")
    print("   1. Unit conversion — longest-match regex, all aliases work")
    print("   2. Web search — DuckDuckGo retry + Wikipedia + RSS fallback")
    print("   3. Input validation — corrected regex char class")
    print("   4. Document QA — semantic search first, keyword fallback")
    print("   5. Routing — explicit commands → convert → math → docs → search")
    print()
    app.run(host='0.0.0.0', port=port, debug=False)